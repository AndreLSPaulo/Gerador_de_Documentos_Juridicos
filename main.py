import streamlit as st
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment
from openpyxl.utils import range_boundaries
from io import BytesIO
import os
import base64
from datetime import datetime

# Caminho da logomarca (opcional)
logo_path = "MP.png"

def get_image_base64(file_path):
    if not os.path.exists(file_path):
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

def formatar_data_extenso(data_str, cidade, uf):
    try:
        meses_pt = {
            1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
            7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
        }
        data = datetime.strptime(data_str, "%d/%m/%Y")
        dia = data.day
        mes = meses_pt[data.month]
        ano = data.year
        return f"{cidade} - {uf}, {dia:02d} de {mes} de {ano}."
    except ValueError:
        return f"{cidade} - {uf}"

def get_top_left_of_merged_cell(ws, cell_coordinate):
    for merged_range in ws.merged_cells.ranges:
        if cell_coordinate in merged_range:
            min_col, min_row, _, _ = range_boundaries(str(merged_range))
            return ws.cell(row=min_row, column=min_col)
    return ws[cell_coordinate]

image_base64 = get_image_base64(logo_path)
if image_base64:
    st.markdown(
        f"""
        <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 20px;">
            <img src="data:image/png;base64,{image_base64}" alt="Logomarca" style="width: 300px;">
        </div>
        """,
        unsafe_allow_html=True,
    )

st.title("Gerador de Documentos Jurídicos")
st.subheader("Preencha os dados do cliente")

def carregar_variaveis():
    dados = {
        "CLIENTE": st.text_input("Nome do Cliente"),
        "ESTADO_CIVIL": st.text_input("Estado Civil"),
        "DATA_NASC": st.text_input("Data de Nascimento"),
        "PROFISSAO": st.text_input("Profissão"),
        "RG": st.text_input("RG"),
        "ORGAO_EXPEDIDOR": st.text_input("Órgão Expedidor"),
        "CPF": st.text_input("CPF"),
        "ENDERECO": st.text_input("Endereço"),
        "Nº": st.text_input("Número"),
        "BAIRRO": st.text_input("Bairro"),
        "COMPLEMENTO": st.text_input("Complemento"),
        "CEP": st.text_input("CEP")
    }

    col1, col2 = st.columns([1, 1])
    with col2:
        cidade_editavel = st.checkbox("Editar cidade manualmente?", key="editar_cidade")
    with col1:
        cidade_padrao = "Manaus"
        dados["CIDADE"] = st.text_input("Cidade", value=cidade_padrao if not cidade_editavel else "", key="cidade_input")

    col3, col4 = st.columns([1, 1])
    with col4:
        comarca_editavel = st.checkbox("Editar comarca manualmente?", key="editar_comarca")
    with col3:
        comarca_padrao = dados["CIDADE"]
        dados["COMARCA"] = st.text_input("Comarca", value=comarca_padrao if not comarca_editavel else "", key="comarca_input")

    dados.update({
        "UF": st.text_input("UF"),
        "DATA": st.text_input("Data"),
        "EMAIL": st.text_input("Email"),
        "WHATSAPP": st.text_input("WhatsApp"),
        "SENHA_GOV": st.text_input("Senha GOV"),
        "TELEFONE2": st.text_input("Telefone 2"),
        "TELEFONE3": st.text_input("Telefone 3"),
        "INDICACAO_CLIENTE": st.text_input("Indicação do Cliente"),
        "PARCERIA_ADVOGADO": st.text_input("Parceria Advogado"),
        "ATENDENTE": st.text_input("Atendente"),
        "TESTEMUNHA1": st.text_input("TESTEMUNHA 1 - Nome"),
        "TESTEMUNHA1_CPF": st.text_input("TESTEMUNHA 1 - CPF"),
        "TESTEMUNHA1_RG": st.text_input("TESTEMUNHA 1 - RG"),
        "TESTEMUNHA1_ORGAO_EXP": st.text_input("TESTEMUNHA 1 - Órgão Expedidor"),
        "TESTEMUNHA1_END": st.text_input("TESTEMUNHA 1 - Endereço"),
        "TESTEMUNHA2": st.text_input("TESTEMUNHA 2 - Nome"),
        "TESTEMUNHA2_CPF": st.text_input("TESTEMUNHA 2 - CPF"),
        "TESTEMUNHA2_RG": st.text_input("TESTEMUNHA 2 - RG"),
        "TESTEMUNHA2_ORGAO_EXP": st.text_input("TESTEMUNHA 2 - Órgão Expedidor"),
        "TESTEMUNHA2_END": st.text_input("TESTEMUNHA 2 - Endereço"),
        "DECLARANTE": st.text_input("Nome do Declarante"),
        "DECLARANTE_ESTADO_CIVIL": st.text_input("Estado Civil do Declarante"),
        "DECLARANTE_PROFISSAO": st.text_input("Profissão do Declarante")
    })
    return dados

dados = carregar_variaveis()

# Pergunta: Cliente alfabetizado?
st.subheader("Cliente alfabetizado(a)?")

opcao_alfabetizado = st.radio(
    "Selecione uma opção:",
    ["Sim", "Não"],
    index=0,
    horizontal=True
)

# Define os modelos de acordo com a escolha
if opcao_alfabetizado == "Sim":
    modelos_arquivo = {
        "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS": "contratos_cadastro/CONTR.PREST.SERV.ADV.xlsx",
        "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA": "contratos_cadastro/DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA.xlsx",
        "DECLARAÇÃO DE VIDA E RESIDÊNCIA": "contratos_cadastro/DECLAR.VIDA.RESIDÊNCIA.xlsx",
        "PROCURAÇÃO AD JUDICIA ET EXTRA": "contratos_cadastro/PROCURAÇÃO AD JUDICIA ET EXTRA.xlsx"
    }
else:
    modelos_arquivo = {
        "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS_ANALFABETO(A)": "contratos_cadastro/CONTR.PREST.SERV.ADV_ANALFABETO(A).xlsx",
        "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA_ANALFABETO(A)": "contratos_cadastro/DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA_ANALFABETO(A).xlsx",
        "DECLARAÇÃO DE VIDA E RESIDÊNCIA_ANALFABETO(A)": "contratos_cadastro/DECLARAÇÃO_DE_VIDA_E_RESIDÊNCIA_ANALFABETO(A).xlsx",
        "PROCURAÇÃO AD JUDICIA ET EXTRA_ANALFABETO(A)": "contratos_cadastro/PROCURAÇÃO AD JUDICIA ET EXTRA_ANALFABETO(A).xlsx"
    }

    # Campos adicionais se o cliente NÃO for alfabetizado
    dados_rogratario = {}
    if opcao_alfabetizado == "Não":
        st.subheader("Preencha os dados do Rogatário (representante do declarante analfabeto)")
        dados_rogratario["ROGATARIO_NOME"] = st.text_input("Nome do Rogatário")
        dados_rogratario["ROGATARIO_RG"] = st.text_input("RG do Rogatário")
        dados_rogratario["ROGATARIO_CPF"] = st.text_input("CPF do Rogatário")
        dados_rogratario["ROGATARIO_END"] = st.text_input("Endereço do Rogatário")

        # Atualiza no dicionário principal
        dados.update(dados_rogratario)
modelo_escolhido = st.selectbox("Modelo disponível", list(modelos_arquivo.keys()))

if st.button("Gerar documento preenchido"):
    caminho_excel = modelos_arquivo.get(modelo_escolhido)
    if caminho_excel and os.path.exists(caminho_excel):
        wb = load_workbook(caminho_excel)
        ws = wb.active

        fonte_padrao = Font(name='Arial', size=11)
        fonte_sublinhada = Font(name='Arial', size=11, underline='single')
        alinhamento_esquerda = Alignment(horizontal='left')
        alinhamento_justificado_topo = Alignment(horizontal='justify', vertical='top', wrap_text=True)
        alinhamento_direita_topo = Alignment(horizontal='right', vertical='top')

        if modelo_escolhido == "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA":
            texto = (
                f"Eu, {dados['CLIENTE']}, brasileiro (a), estado civil {dados['ESTADO_CIVIL']}, Profissão {dados['PROFISSAO']}, "
                f"RG {dados['RG']} SSP {dados['UF']}. Inscrito no CPF sob o nº {dados['CPF']} residente e domiciliado nesta cidade de "
                f"{dados['CIDADE']}/{dados['UF']} na {dados['ENDERECO']} - {dados['COMPLEMENTO']} Nº {dados['Nº']}, Bairro: {dados['BAIRRO']}, CEP:{dados['CEP']}"
            )
            ws["A8"] = texto
            ws["A8"].font = fonte_padrao
            ws["A8"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 30

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["D25"] = data_ext
            ws["D25"].font = fonte_sublinhada
            ws["D25"].alignment = alinhamento_direita_topo

        elif modelo_escolhido == "DECLARAÇÃO DE VIDA E RESIDÊNCIA":
            texto = (
                f"Eu, {dados['DECLARANTE']}. Brasileiro (a), estado civil {dados['DECLARANTE_ESTADO_CIVIL']}, Profissão {dados['DECLARANTE_PROFISSAO']}.\n"
                f"DECLARO para os devidos fins de comprovação de residência, sob as penas da lei (art. 2º da lei 7.115/83), que {dados['CLIENTE']}, "
                f"Brasileiro (a), Estado Civil {dados['ESTADO_CIVIL']}, Profissão: {dados['PROFISSAO']}, portador(a), do RG: {dados['RG']} e CPF: {dados['CPF']}, "
                f"é residente e domiciliado na {dados['ENDERECO']} - Nº {dados['Nº']} - {dados['COMPLEMENTO']}. Bairro: {dados['BAIRRO']}. "
                f"CEP: {dados['CEP']}. Cidade: {dados['CIDADE']}, UF: {dados['UF']}."
            )
            ws["A8"] = texto
            ws["A8"].font = fonte_padrao
            ws["A8"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 60

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["I17"] = data_ext
            ws["I17"].font = fonte_padrao
            ws["I17"].alignment = alinhamento_direita_topo

        elif modelo_escolhido == "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS":
            preenchimentos = {
                "C8": "CLIENTE",
                "B10": "ESTADO_CIVIL",
                "G10": "PROFISSAO",
                "B12": "RG",
                "F12": "ORGAO_EXPEDIDOR",
                "H12": "CPF",
                "B14": "ENDERECO",
                "B16": "CEP",
                "H14": "COMPLEMENTO",
                "F1": "INDICACAO_CLIENTE",
                "F2": "PARCERIA_ADVOGADO",
                "F3": "ATENDENTE",
                "E16": "CIDADE",
                "I16": "UF",
                "B18": "EMAIL",
                "B20": "WHATSAPP",
                "E20": "TELEFONE2",
                "H20": "TELEFONE3",
                "G68": "CIDADE",
                "B21": "SENHA_GOV",
                "B84": "TESTEMUNHA1",
                "B85": "TESTEMUNHA1_CPF",
                "B86": "TESTEMUNHA1_RG",
                "B87": "TESTEMUNHA1_END",
                "G84": "TESTEMUNHA2",
                "G85": "TESTEMUNHA2_CPF",
                "G86": "TESTEMUNHA2_RG",
                "G87": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

            clausula = (
                "I) 30% (trinta por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
                if dados.get("CIDADE", "").lower() == "manaus" else
                "I) 35% (trinta e cinco por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
            )
            ws["A29"] = clausula
            ws["A29"].font = fonte_padrao
            ws["A29"].alignment = alinhamento_justificado_topo

            for linha in ["A31", "A34", "A61", "A63", "A65", "B84", "B85", "B86", "B87", "G84", "G85", "G86", "G87"]:
                ws[linha].font = fonte_padrao
                ws[linha].alignment = alinhamento_justificado_topo

            clausula_comarca = f"CLÁUSULA 17ª - As  partes  contratantes  elegem  o  foro  da  Comarca  de {dados.get('COMARCA', '')} para dirimir quaisquer controvérsias oriundas do presente contrato."
            ws["A68"] = clausula_comarca
            ws["A68"].font = fonte_padrao
            ws["A68"].alignment = alinhamento_justificado_topo

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["F72"] = data_ext
            ws["F72"].font = fonte_padrao
            ws["F72"].alignment = alinhamento_direita_topo

            ws["A76"] = dados.get("CLIENTE", "")
            ws["A76"].font = fonte_padrao
            ws["A76"].alignment = alinhamento_justificado_topo
        
        elif modelo_escolhido == "PROCURAÇÃO AD JUDICIA ET EXTRA":
            preenchimentos = {
                "C3": "CLIENTE",
                "C4": "ESTADO_CIVIL",
                "G4": "PROFISSAO",
                "B5": "RG",
                "G5": "ORGAO_EXPEDIDOR",
                "B6": "CPF",
                "B9": "EMAIL"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda
                
                 # Construir endereço completo manualmente e aplicar na célula A8
                endereco_formatado = (
                    f"{dados.get('ENDERECO', '')}, Nº {dados.get('Nº', '')}, "
                    f"{dados.get('BAIRRO', '')}, CEP: {dados.get('CEP', '')}, {dados.get('COMPLEMENTO', '')}"
                )
                ws["A8"] = endereco_formatado
                ws["A8"].font = fonte_padrao
                ws["A8"].alignment = alinhamento_justificado_topo

                data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
                ws["I17"] = data_ext
                ws["I17"].font = fonte_padrao
                ws["I17"].alignment = alinhamento_direita_topo

        elif modelo_escolhido == "DECLARAÇÃO DE VIDA E RESIDÊNCIA_ANALFABETO(A)":
                texto = (
                    f"Eu, {dados['DECLARANTE']}. Brasileiro (a), estado civil {dados['DECLARANTE_ESTADO_CIVIL']}, Profissão {dados['DECLARANTE_PROFISSAO']}.\n"
                    f"DECLARO para os devidos fins de comprovação de residência, sob as penas da lei (art. 2º da lei 7.115/83), que {dados['CLIENTE']}, "
                    f"Brasileiro (a), Estado Civil {dados['ESTADO_CIVIL']}, Profissão: {dados['PROFISSAO']}, portador(a), do RG: {dados['RG']} e CPF: {dados['CPF']}, "
                    f"é residente e domiciliado na {dados['ENDERECO']} - Nº {dados['Nº']} - {dados['COMPLEMENTO']}. Bairro: {dados['BAIRRO']}. "
                    f"CEP: {dados['CEP']}. Cidade: {dados['CIDADE']}, UF: {dados['UF']}."
                )
                ws["A6"] = texto
                ws["A6"].font = fonte_padrao
                ws["A6"].alignment = alinhamento_justificado_topo
                ws.row_dimensions[8].height = 60

                data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
                ws["I18"] = data_ext
                ws["I18"].font = fonte_padrao
                ws["I18"].alignment = alinhamento_direita_topo

                preenchimentos = {
                    "B26": "ROGATARIO_NOME",
                    "B27": "ROGATARIO_RG",
                    "B28": "ROGATARIO_CPF",
                    "B29": "ROGATARIO_END",
                    "B32": "TESTEMUNHA1",
                    "B34": "TESTEMUNHA1_CPF",
                    "B33": "TESTEMUNHA1_RG",
                    "B35": "TESTEMUNHA1_END",
                    "B38": "TESTEMUNHA2",
                    "B40": "TESTEMUNHA2_CPF",
                    "B39": "TESTEMUNHA2_RG",
                    "B41": "TESTEMUNHA2_END"
                }

                for celula, campo in preenchimentos.items():
                    valor = dados.get(campo, "")
                    target_cell = get_top_left_of_merged_cell(ws, celula)
                    target_cell.value = valor
                    target_cell.font = fonte_padrao
                    target_cell.alignment = alinhamento_esquerda

        elif modelo_escolhido == "PROCURAÇÃO AD JUDICIA ET EXTRA_ANALFABETO(A)":
                preenchimentos = {
                    "C3": "CLIENTE",
                    "C4": "ESTADO_CIVIL",
                    "F4": "PROFISSAO",
                    "B5": "RG",
                    "G5": "ORGAO_EXPEDIDOR",
                    "B6": "CPF",
                    "B9": "EMAIL"
                }

                for celula, campo in preenchimentos.items():
                    valor = dados.get(campo, "")
                    target_cell = get_top_left_of_merged_cell(ws, celula)
                    target_cell.value = valor
                    target_cell.font = fonte_padrao
                    target_cell.alignment = alinhamento_esquerda

                endereco_formatado = (
                    f"{dados.get('ENDERECO', '')}, Nº {dados.get('Nº', '')}, "
                    f"{dados.get('BAIRRO', '')}, CEP: {dados.get('CEP', '')}, {dados.get('COMPLEMENTO', '')}"
                )
                ws["A8"] = endereco_formatado
                ws["A8"].font = fonte_padrao
                ws["A8"].alignment = alinhamento_justificado_topo

                data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
                ws["K17"] = data_ext
                ws["K17"].font = fonte_padrao
                ws["K17"].alignment = alinhamento_direita_topo

                # Preencher a célula E23 com: "NOME, CPF: xxx.xxx.xxx-xx"
                nome_roga = dados.get("ROGATARIO_NOME", "").strip()
                cpf_roga = dados.get("ROGATARIO_CPF", "").strip()
                valor_e23 = f"{nome_roga}, CPF: {cpf_roga}" if nome_roga else ""
                target_cell = get_top_left_of_merged_cell(ws, "E23")
                target_cell.value = valor_e23
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

                # Preencher as demais células
                preenchimentos = {
                    "B26": "TESTEMUNHA1",
                    "B28": "TESTEMUNHA1_CPF",
                    "B27": "TESTEMUNHA1_END",
                    "G26": "TESTEMUNHA2",
                    "G28": "TESTEMUNHA2_CPF",
                    "G27": "TESTEMUNHA2_END"
                }

                for celula, campo in preenchimentos.items():
                    valor = dados.get(campo, "")
                    target_cell = get_top_left_of_merged_cell(ws, celula)
                    target_cell.value = valor
                    target_cell.font = fonte_padrao
                    target_cell.alignment = alinhamento_esquerda

        elif modelo_escolhido == "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA_ANALFABETO(A)":
            texto = (
                f"Eu, {dados['CLIENTE']}, brasileiro (a), estado civil {dados['ESTADO_CIVIL']}, Profissão {dados['PROFISSAO']}, "
                f"RG {dados['RG']} SSP {dados['UF']}. Inscrito no CPF sob o nº {dados['CPF']} residente e domiciliado nesta cidade de "
                f"{dados['CIDADE']}/{dados['UF']} na {dados['ENDERECO']} - {dados['COMPLEMENTO']} Nº {dados['Nº']}, Bairro: {dados['BAIRRO']}, CEP:{dados['CEP']}"
            )
            ws["A6"] = texto
            ws["A6"].font = fonte_padrao
            ws["A6"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 30

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["K20"] = data_ext
            ws["K20"].font = fonte_sublinhada
            ws["K20"].alignment = alinhamento_direita_topo

            preenchimentos = {
                "B24": "ROGATARIO_NOME",
                "B25": "ROGATARIO_RG",
                "B26": "ROGATARIO_CPF",
                "B27": "ROGATARIO_END",
                "B30": "TESTEMUNHA1",
                "B31": "TESTEMUNHA1_RG",
                "B32": "TESTEMUNHA1_CPF",
                "B33": "TESTEMUNHA1_END",
                "B36": "TESTEMUNHA2",
                "B37": "TESTEMUNHA2_RG",
                "B38": "TESTEMUNHA2_CPF",
                "B39": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

        elif modelo_escolhido == "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS_ANALFABETO(A)":
            preenchimentos = {
                "C8": "CLIENTE",
                "B10": "ESTADO_CIVIL",
                "G10": "PROFISSAO",
                "B12": "RG",
                "F12": "ORGAO_EXPEDIDOR",
                "H12": "CPF",
                "B14": "ENDERECO",
                "B16": "CEP",
                "G14": "COMPLEMENTO",
                "F1": "INDICACAO_CLIENTE",
                "F2": "PARCERIA_ADVOGADO",
                "F3": "ATENDENTE",
                "E16": "CIDADE",
                "I16": "UF",
                "B18": "EMAIL",
                "B20": "WHATSAPP",
                "E20": "TELEFONE2",
                "H20": "TELEFONE3",
                "G68": "CIDADE",
                "B21": "SENHA_GOV",
		        "B79": "ROGATARIO_NOME",
                "B80": "ROGATARIO_RG",
                "B81": "ROGATARIO_CPF",
                "B82": "ROGATARIO_END",		    
		        "B85": "TESTEMUNHA1",
                "B87": "TESTEMUNHA1_CPF",
                "B86": "TESTEMUNHA1_RG",
                "B88": "TESTEMUNHA1_END",
                "B91": "TESTEMUNHA2",
                "B93": "TESTEMUNHA2_CPF",
                "B92": "TESTEMUNHA2_RG",
                "B94": "TESTEMUNHA2_END",
                "B102": "TESTEMUNHA1",
                "B103": "TESTEMUNHA1_CPF",
                "B104": "TESTEMUNHA1_RG",
		        "B105": "TESTEMUNHA1_END",
                "G102": "TESTEMUNHA2",
                "G103": "TESTEMUNHA2_CPF",
                "G104": "TESTEMUNHA2_RG",
                "G105": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

            clausula = (
                "I) 30% (trinta por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
                if dados.get("CIDADE", "").lower() == "manaus" else
                "I) 35% (trinta e cinco por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
            )
            ws["A29"] = clausula
            ws["A29"].font = fonte_padrao
            ws["A29"].alignment = alinhamento_justificado_topo

            for linha in ["A31", "A34", "A58", "A64"]:
                ws[linha].font = fonte_padrao
                ws[linha].alignment = alinhamento_justificado_topo

            clausula_comarca = f"CLÁUSULA 17ª - As  partes  contratantes  elegem  o  foro  da  Comarca  de {dados.get('COMARCA', '')} para dirimir quaisquer controvérsias oriundas do presente contrato."
            ws["A67"] = clausula_comarca
            ws["A67"].font = fonte_padrao
            ws["A67"].alignment = alinhamento_justificado_topo

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["I71"] = data_ext
            ws["I71"].font = fonte_padrao
            ws["I71"].alignment = alinhamento_direita_topo

            ws["A74"] = dados.get("CLIENTE", "")
            ws["A74"].font = fonte_padrao
            ws["A74"].alignment = alinhamento_justificado_topo
        
       

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        st.download_button(
            label="📥 Baixar documento preenchido",
            data=output,
            file_name=f"{modelo_escolhido.replace(' ', '_')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        st.success("✅ Planilha preenchida com sucesso!")
    else:
        st.error("❌ Arquivo de modelo não encontrado.")

# ==============================
# RECIBOS DE SERVIÇOS JURÍDICOS
# ==============================
from typing import List
from docx import Document
from io import BytesIO

# ---------- Utilidades p/ moeda PT-BR ----------
def parse_valor_brl(s: str) -> float:
    """Converte '1.234,56' | '1234,56' | '1234.56' -> 1234.56 (float)."""
    if not s:
        return 0.0
    s = s.strip().replace("R$", "").replace(" ", "")
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return 0.0

def formatar_brl(v: float) -> str:
    """Formata 1234.56 -> '1.234,56'."""
    return f"{v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def _extenso_0_999(n: int) -> str:
    unidades = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
    dez_a_dezenove = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
    dezenas = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
    centenas = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
    if n == 0: return ""
    if n == 100: return "cem"
    c = n // 100; d = (n % 100) // 10; u = n % 10
    partes = []
    if c: partes.append(centenas[c])
    if d == 1:
        partes.append(dez_a_dezenove[u])
    else:
        if d: partes.append(dezenas[d])
        if u: partes.append(unidades[u])
    saida = ""
    for p in partes:
        saida = p if not saida else f"{saida} e {p}"
    return saida

def _bloco_extenso(n: int, singular: str, plural: str) -> str:
    if n == 0: return ""
    if n == 1: return f"um {singular}"
    return f"{_extenso_0_999(n)} {plural}"

def numero_para_moeda_ptbr(valor: float) -> str:
    """
    1234.56 -> 'Mil duzentos e trinta e quatro reais e cinquenta e seis centavos'
    (primeira letra maiúscula).
    """
    if valor < 0:
        frase = "menos " + numero_para_moeda_ptbr(-valor)
        return frase[0].upper() + frase[1:]

    inteiro = int(valor)
    centavos = int(round((valor - inteiro) * 100))
    bilhoes = inteiro // 1_000_000_000
    resto = inteiro % 1_000_000_000
    milhoes = resto // 1_000_000
    resto %= 1_000_000
    milhares = resto // 1_000
    centenas = resto % 1_000

    partes = []
    if bilhoes: partes.append(_bloco_extenso(bilhoes, "bilhão", "bilhões"))
    if milhoes: partes.append(_bloco_extenso(milhoes, "milhão", "milhões"))
    if milhares:
        partes.append("mil" if milhares == 1 else f"{_extenso_0_999(milhares)} mil")
    if centenas: partes.append(_extenso_0_999(centenas))

    partes_reais = "zero" if not partes else " ".join(partes).replace("mil e ", "mil ")
    sufx_reais = "real" if inteiro == 1 else "reais"
    frase = f"{partes_reais} {sufx_reais}"

    if centavos:
        ext_cent = _extenso_0_999(centavos)
        sufx_cent = "centavo" if centavos == 1 else "centavos"
        frase += f" e {ext_cent} {sufx_cent}"

    # inicial maiúscula
    return frase[0].upper() + frase[1:] if frase else frase
# ---------- fim utilidades moeda ----------

st.divider()
st.subheader("Recibo de serviços jurídicos")

opcoes_recibo = [
    "CONSULTORIA JURÍDICA",
    "SEGUNDA PARCELA DOS HONORÁRIOS",
    "PGTO PARCELADO",
    "PGTO ÚNICO – PROCESSOS ADMINISTRATIVOS JUNTO AO INSS",
    "PGTO – PRIMEIRA PARCELA DOS HONORÁRIOS PARCIAIS DO PROCESSO",
    "PROCESSO DE ANÁLISE DE DESCONTOS DE ENTIDADES ASSOCIATIVAS",
    "QUANDO AUXILIAMOS A FAZER A RECLAMAÇÃO",
    "RECLAMAÇÃO DO CARTÃO CRÉDITO/DÉBITO",
    "RESGATE TIT. CAPITALIZAÇÃO",
    "SIMPLES",
]

# Caminho base do modelo .docx (robusto)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH_1 = os.path.join(BASE_DIR, "contratos_cadastro", "Recibo de servicos juridicos.docx")
DOCX_PATH_2 = os.path.join(BASE_DIR, "Recibo de servicos juridicos.docx")  # fallback
TEMPLATE_DOCX = DOCX_PATH_1 if os.path.exists(DOCX_PATH_1) else DOCX_PATH_2

selecionada = st.selectbox("Selecione o tipo de recibo:", opcoes_recibo, index=0, key="tipo_recibo")

# ---- Campos extras apenas para "CONSULTORIA JURÍDICA"
valor_str = ""
hora = ""
valor_extenso_auto = ""
valor_formatado = ""

if selecionada == "CONSULTORIA JURÍDICA":
    colv1, colv2 = st.columns([1, 1])
    with colv1:
        valor_str = st.text_input("VALOR (R$) — {VALOR}", placeholder="300,00")
    with colv2:
        hora = st.text_input("HORA — {HORA}", placeholder="11:00")

    valor_float = parse_valor_brl(valor_str)
    valor_formatado = formatar_brl(valor_float) if valor_str else ""
    valor_extenso_auto = numero_para_moeda_ptbr(valor_float)  # já vem com inicial maiúscula

    editar_extenso = st.checkbox("Editar valor por extenso manualmente?", value=False)
    st.text_input(
        "VALOR_EXTENSO — ({VALOR_EXTENSO})",
        value=valor_extenso_auto,
        disabled=not editar_extenso,
        key="valor_extenso"
    )

# ---- Placeholders
data_extenso = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))

valor_extenso_final = st.session_state.get("valor_extenso") or valor_extenso_auto  # já capitalizado na função

placeholders = {
    "{CLIENTE}": dados.get("CLIENTE", ""),
    "{CPF}": dados.get("CPF", ""),
    "{VALOR}": valor_formatado if selecionada == "CONSULTORIA JURÍDICA" else "",
    "{VALOR_EXTENSO}": valor_extenso_final if selecionada == "CONSULTORIA JURÍDICA" else "",
    "{DATA}": dados.get("DATA", ""),
    "{HORA}": hora if selecionada == "CONSULTORIA JURÍDICA" else "",
    "{CIDADE}": dados.get("CIDADE", ""),
    "{UF}": dados.get("UF", ""),
    "{DATA em extenso}": data_extenso,
    "{DATA_EXTENSO}": data_extenso,
}

# ---- Texto padrão do recibo "CONSULTORIA JURÍDICA"
texto_base_consultoria = (
    "Recebi de {CLIENTE}, portador (a) do CPF {CPF}, a importância de R$ {VALOR} ({VALOR_EXTENSO}), "
    "face à consultoria jurídica realizada no dia {DATA}, às {HORA} horas, qual dou plena quitação.\n\n"
    
)

def preencher_texto(modelo: str, mapping: dict) -> str:
    out = modelo
    for k, v in mapping.items():
        out = out.replace(k, v if v is not None else "")
    return out

# ---- Pré-visualização (preenche apenas para a opção "CONSULTORIA JURÍDICA")
preview_text = ""
preview_editado = ""
if selecionada == "CONSULTORIA JURÍDICA":
    preview_text = preencher_texto(texto_base_consultoria, placeholders)

    st.markdown("**Pré-visualização (editável):**")
    preview_editado = st.text_area(
        "Você pode ajustar o texto antes de gerar o arquivo:",
        value=preview_text,
        height=220
    )

# ---- Utilidades de manipulação do .docx
def replace_in_paragraph(paragraph, mapping: dict):
    novo = preencher_texto(paragraph.text, mapping)
    if novo != paragraph.text:
        paragraph.text = novo  # substitui placeholders; perde estilos de runs

def replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)

def _inserir_3_linhas_apos_titulo(doc: Document, titulo_ref: str, linhas: List[str]):
    """
    Procura o parágrafo com 'titulo_ref' e insere os itens de 'linhas'
    exatamente 3 linhas abaixo. Cria parágrafos vazios se necessário.
    """
    titulo_ref_upper = titulo_ref.upper()
    for i, p in enumerate(doc.paragraphs):
        if titulo_ref_upper in (p.text or "").upper():
            insert_index = i + 3
            # garanta espaço até a posição de inserção
            while len(doc.paragraphs) <= insert_index:
                doc.add_paragraph("")
            # insere as linhas na ordem, cada uma como novo parágrafo
            for linha in linhas:
                doc.paragraphs[insert_index].insert_paragraph_before(linha)
                insert_index += 1
            break

def render_docx_from_template(
    template_path: str,
    mapping: dict,
    linhas_consultoria: List[str] | None = None,
    data_extenso_str: str = ""
) -> BytesIO:
    doc = Document(template_path)

    # 1) substitui placeholders existentes
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for t in doc.tables:
        replace_in_table(t, mapping)

    # 2) insere o texto 3 linhas abaixo do título
    if linhas_consultoria:
        _inserir_3_linhas_apos_titulo(
            doc,
            "RECIBO DE PAGAMENTO",
            linhas_consultoria
        )

    # 3) insere {DATA em extenso} 2 linhas acima de "MARCELA DA SILVA PAULO" (à direita)
    if data_extenso_str:
        alvo_upper = "MARCELA DA SILVA PAULO"
        for i, p in enumerate(doc.paragraphs):
            if alvo_upper in (p.text or "").upper():
                insert_index = max(i - 2, 0)
                novo = doc.paragraphs[insert_index].insert_paragraph_before(data_extenso_str)
                novo.alignment = 2  # right
                break

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio



# ---- Botões de geração/Download
colg1, colg2 = st.columns([1, 1])
with colg1:
    gerar = st.button("🔄 Gerar arquivo (.docx)")
with colg2:
    st.write("")

if gerar:
    if not os.path.exists(TEMPLATE_DOCX):
        st.error(
            "❌ Arquivo base não encontrado.\n"
            f"Verifique:\n- {DOCX_PATH_1}\n- {DOCX_PATH_2}"
        )
    else:
        # monta o texto final da consultoria a partir da PRÉ-VISUALIZAÇÃO editável (se houver)
        if selecionada == "CONSULTORIA JURÍDICA":
            texto_final_consultoria = preview_editado or preencher_texto(texto_base_consultoria, placeholders)
        else:
            texto_final_consultoria = preencher_texto(texto_base_consultoria, placeholders)

        linhas_para_inserir = [linha for linha in texto_final_consultoria.split("\n")]

        buffer_docx = render_docx_from_template(
    TEMPLATE_DOCX,
    placeholders,
    linhas_consultoria=linhas_para_inserir,
    data_extenso_str=placeholders.get("{DATA em extenso}", "")
)

        st.success("✅ Arquivo gerado. Clique para baixar:")
        st.download_button(
            label="📥 Baixar Recibo (.docx)",
            data=buffer_docx,
            file_name="Recibo_de_servicos_juridicos.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
